"""
Generate Word (DOCX) project report for VastraVista with code excerpts and explanations.
"""

import os
from pathlib import Path
import ast
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

BASE_DIR = Path(__file__).resolve().parents[2]
APP_DIR = BASE_DIR
OUTPUT_DIR = APP_DIR / 'data' / 'reports'


def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def ensure_code_style(doc: Document):
    styles = doc.styles
    if 'Code' not in [s.name for s in styles]:
        style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'Courier New'
        font.size = Pt(10)
    return styles['Code']


def read_text(path: Path) -> str:
    try:
        return path.read_text(encoding='utf-8')
    except Exception:
        return ''


def get_function_source(file_path: Path, func_name: str) -> str:
    try:
        src = read_text(file_path)
        tree = ast.parse(src)
        for node in ast.walk(tree):
            if isinstance(node, ast.FunctionDef) and node.name == func_name:
                segment = ast.get_source_segment(src, node)
                return segment or ''
    except Exception:
        pass
    return ''


def get_js_function_source(file_path: Path, func_name: str) -> str:
    src = read_text(file_path)
    if not src:
        return ''
    sigs = [f"function {func_name}", f"const {func_name} =", f"let {func_name} =", f"var {func_name} ="]
    start = -1
    for s in sigs:
        start = src.find(s)
        if start != -1:
            break
    if start == -1:
        return ''
    brace_start = src.find('{', start)
    if brace_start == -1:
        return ''
    depth = 0
    end = brace_start
    for i, ch in enumerate(src[brace_start:], brace_start):
        if ch == '{':
            depth += 1
        elif ch == '}':
            depth -= 1
            if depth == 0:
                end = i + 1
                break
    segment = src[start:end]
    return segment


def add_code_block(doc: Document, title: str, file_path: Path, func_name: str, explanation: str):
    add_heading(doc, title, level=3)
    p = doc.add_paragraph(explanation)
    p.style = doc.styles['Normal']
    style = ensure_code_style(doc)
    code = ''
    if file_path.suffix == '.py':
        code = get_function_source(file_path, func_name)
    elif file_path.suffix == '.js':
        code = get_js_function_source(file_path, func_name)
    if not code:
        code = f"[Could not extract function '{func_name}' from {file_path.name}]"
    cp = doc.add_paragraph(code)
    cp.style = style


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    out_path = OUTPUT_DIR / 'VastraVista_Project_Report.docx'

    doc = Document()

    # Title
    add_heading(doc, 'VastraVista – Project Report', level=0)
    doc.add_paragraph('AI Fashion Insights & Color Analysis (Word Report with Code Excerpts)')

    # Overview
    add_heading(doc, 'Overview', level=1)
    doc.add_paragraph(
        'VastraVista analyzes user photos to detect skin tone, gender, age, and recommends complementary '
        'clothing colors. It provides AI-generated insights, multi-face analysis, and outfit color feedback.'
    )

    # Architecture
    add_heading(doc, 'Architecture', level=1)
    doc.add_paragraph('Backend services in Flask; frontend in JavaScript; MediaPipe for face mesh; Delta‑E CIE2000 for color distance.')
    doc.add_paragraph('Key modules: app/api/auth.py, app/services/ar_styling_service.py, app/services/color_analyzer.py, app/services/ai_stylist.py, static/js/app.js')

    # Endpoints
    add_heading(doc, 'Endpoints', level=1)
    doc.add_paragraph('/api/analyze – core analysis upload and response')
    doc.add_paragraph('/uploads/<filename> – serve uploaded images')
    doc.add_paragraph('/api/v2/ai-fashion-advice – AI advice from existing analysis')
    doc.add_paragraph('/api/v2/monk-scale-info – Monk scale data and palettes')

    # Code Excerpts
    add_heading(doc, 'Core Code & Explanations', level=1)

    add_code_block(
        doc,
        'Analysis Endpoint (/api/analyze)',
        APP_DIR / 'app' / 'api' / 'auth.py',
        'api_analyze',
        'Handles the upload, preprocessing, detection of skin/gender/age, color recommendations, '
        'per-face clothing color extraction, outfit feedback, and AI independent analysis/verification.'
    )

    add_code_block(
        doc,
        'Dominant Clothing Color Extraction (full image)',
        APP_DIR / 'app' / 'services' / 'ar_styling_service.py',
        'extract_dominant_clothing_color',
        'Uses MediaPipe face mesh to mask collar/shoulders, clusters pixels via k-means, maps to nearest fashion color.'
    )

    add_code_block(
        doc,
        'Per-Face Clothing Color Extraction (bbox)',
        APP_DIR / 'app' / 'services' / 'ar_styling_service.py',
        'extract_clothing_color_for_bbox',
        'Runs the same extraction within a face ROI to support multi-face images, producing person-specific feedback.'
    )

    add_code_block(
        doc,
        'Fashion Color Scoring (Delta‑E + rules)',
        APP_DIR / 'app' / 'services' / 'color_analyzer.py',
        'find_best_colors',
        'Computes complementary colors using Delta‑E (Lab) plus fashion heuristics for brightness, saturation, and skin category.'
    )

    add_code_block(
        doc,
        'AI Independent Analysis',
        APP_DIR / 'app' / 'services' / 'ai_stylist.py',
        'analyze_image_independently',
        'Generates gender/age/skin tone with a local LLM; used to override or validate technical outputs when confidence is low.'
    )

    add_code_block(
        doc,
        'AI vs Technical Comparison',
        APP_DIR / 'app' / 'services' / 'ai_stylist.py',
        'compare_analyses',
        'Compares AI results with technical analysis to produce an agreement score and highlight differences.'
    )

    add_code_block(
        doc,
        'Frontend Results Rendering',
        APP_DIR / 'static' / 'js' / 'app.js',
        'displayAnalysisResults',
        'Builds dynamic UI: analyzed photo, verification badges, comparison, outfit feedback, and stat cards.'
    )

    add_heading(doc, 'Configuration', level=1)
    doc.add_paragraph('Uploads folder path and report folder ensured in app/config/config.py and app/__init__.py.')

    add_heading(doc, 'Report Path', level=1)
    doc.add_paragraph(str(out_path))

    doc.save(out_path)
    print(f"DOCX report generated: {out_path}")


if __name__ == '__main__':
    main()
